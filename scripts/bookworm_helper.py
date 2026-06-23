#!/usr/bin/env python3
"""Bookworm helper utilities.

Standard-library-only helpers for deterministic book inspection, EPUB image
extraction, and Obsidian vault discovery. LLM summarization is handled by Codex
through the Bookworm skill.
"""

from __future__ import annotations

import argparse
import html
import json
import os
import re
import shutil
import sys
import zipfile
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path, PurePosixPath
from typing import Iterable
import xml.etree.ElementTree as ET


HTML_EXTENSIONS = (".html", ".xhtml", ".htm")
IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg")
REFINE_INPUT_EXTENSIONS = (".md", ".docx", ".pdf", ".pptx")
CHATGPT_CITATION_PATTERN = re.compile(r"cite[^]*")
NUMERIC_CITATION_PATTERN = re.compile(r"(?<!\w)\[(\d{1,4})\]")
MARKDOWN_HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
MARKDOWN_SOURCE_LINK_PATTERN = re.compile(r"(?<!!)\[[^\]]+\]\(\s*https?://[^)\s]+[^)]*\)")
BARE_URL_PATTERN = re.compile(r"(?<!\]\()(?<!\()https?://[^\s<>)\]]+")
FOOTNOTE_REFERENCE_PATTERN = re.compile(r"(?<!\\)\[\^[^\]]+\]")
LABEL_VALUE_PATTERN = re.compile(r"^\*\*(.+?)(?::)?\*\*\s*:?\s*(.+?)\s*$")
MERMAID_COMPACT_INIT = '%%{init: {"flowchart": {"useMaxWidth": false, "nodeSpacing": 20, "rankSpacing": 25}} }%%'


class TextAndImageParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[tuple[str, str]] = []
        self.images: list[str] = []
        self._tag: str | None = None
        self._buf: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        attrs_dict = {k.lower(): v for k, v in attrs if v is not None}
        if tag in {"h1", "h2", "h3", "h4", "p", "li"}:
            self._tag = tag
            self._buf = []
        if tag == "img":
            src = attrs_dict.get("src") or attrs_dict.get("href")
            if src:
                self.images.append(src)

    def handle_data(self, data: str) -> None:
        if self._tag:
            self._buf.append(data)

    def handle_endtag(self, tag: str) -> None:
        if self._tag and tag.lower() == self._tag:
            text = normalize_text(" ".join(self._buf))
            if text:
                self.parts.append((self._tag, text))
            self._tag = None


@dataclass
class BookSource:
    path: Path

    @property
    def is_epub_dir(self) -> bool:
        return self.path.is_dir() and (self.path / "META-INF" / "container.xml").exists()

    @property
    def is_zip(self) -> bool:
        return self.path.is_file() and zipfile.is_zipfile(self.path)

    def names(self) -> list[str]:
        if self.is_epub_dir:
            return [
                str(p.relative_to(self.path)).replace(os.sep, "/")
                for p in self.path.rglob("*")
                if p.is_file()
            ]
        if self.is_zip:
            with zipfile.ZipFile(self.path) as zf:
                return zf.namelist()
        raise ValueError(f"Unsupported EPUB source: {self.path}")

    def read_bytes(self, name: str) -> bytes:
        if self.is_epub_dir:
            return (self.path / name).read_bytes()
        if self.is_zip:
            with zipfile.ZipFile(self.path) as zf:
                return zf.read(name)
        raise ValueError(f"Unsupported EPUB source: {self.path}")


def normalize_text(value: str) -> str:
    value = html.unescape(value)
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def slugify(value: str, fallback: str = "book") -> str:
    value = value.lower()
    translit = {
        "а": "a", "б": "b", "в": "v", "г": "g", "д": "d", "е": "e", "ё": "e",
        "ж": "zh", "з": "z", "и": "i", "й": "y", "к": "k", "л": "l", "м": "m",
        "н": "n", "о": "o", "п": "p", "р": "r", "с": "s", "т": "t", "у": "u",
        "ф": "f", "х": "h", "ц": "ts", "ч": "ch", "ш": "sh", "щ": "sch",
        "ъ": "", "ы": "y", "ь": "", "э": "e", "ю": "yu", "я": "ya",
    }
    value = "".join(translit.get(ch, ch) for ch in value)
    value = re.sub(r"[^a-z0-9]+", "-", value).strip("-")
    return value or fallback


def obsidian_anchor(value: str, fallback: str = "section") -> str:
    """Build a conservative Obsidian-compatible anchor from a heading."""
    value = re.sub(r"[`*_~]", "", value).lower()
    value = re.sub(r"[^\w\s-]", "", value, flags=re.UNICODE)
    value = re.sub(r"[\s-]+", "-", value, flags=re.UNICODE).strip("-")
    return value or fallback


def source_counts(source: str) -> dict[str, int]:
    """Count source-bearing constructs that Refine must not silently lose."""
    markdown_links = MARKDOWN_SOURCE_LINK_PATTERN.findall(source)
    source_without_markdown_links = MARKDOWN_SOURCE_LINK_PATTERN.sub("", source)
    return {
        "citation_markers": len(CHATGPT_CITATION_PATTERN.findall(source)),
        "markdown_links": len(markdown_links),
        "bare_urls": len(BARE_URL_PATTERN.findall(source_without_markdown_links)),
        "footnote_references": len(FOOTNOTE_REFERENCE_PATTERN.findall(source)),
    }


def citation_inventory(source: str) -> list[dict[str, str | int]]:
    """Return raw ChatGPT citations with the paragraph they annotate."""
    entries: list[dict[str, str | int]] = []
    for index, line in enumerate(source.splitlines(), start=1):
        markers = CHATGPT_CITATION_PATTERN.findall(line)
        if not markers:
            continue
        context = line.strip()
        for marker in markers:
            entries.append({"line": index, "marker": marker, "context": context})
    return entries


def resolve_citation_markers(
    source: str,
    verified_sources: dict[str, dict[str, str]],
) -> tuple[str, dict[str, int]]:
    """Replace raw markers only with sources a caller has already verified.

    The resolver deliberately accepts no search terms or guessed URLs: source
    discovery and opening happen in the Refine workflow, while this helper
    deterministically applies those verified choices and reports every gap.
    """
    report = {
        "markers_scanned": 0,
        "verified_title_links_inserted": 0,
        "unresolved": 0,
    }

    def replace(match: re.Match[str]) -> str:
        marker = match.group(0)
        report["markers_scanned"] += 1
        verified = verified_sources.get(marker)
        if not verified or not verified.get("title") or not verified.get("url"):
            report["unresolved"] += 1
            return ""
        report["verified_title_links_inserted"] += 1
        return f" [{verified['title']}]({verified['url']})"

    result = CHATGPT_CITATION_PATTERN.sub(replace, source)
    result = re.sub(r"[ \t]+(\[[^]]+\]\(https?://)", r" \1", result)
    return result, report


def _source_section_bounds(lines: list[str]) -> tuple[int, int] | None:
    for index, line in enumerate(lines):
        heading = MARKDOWN_HEADING_PATTERN.match(line)
        if not heading or heading.group(2).strip().casefold() not in {"sources", "источники"}:
            continue
        level = len(heading.group(1))
        end = index + 1
        while end < len(lines):
            next_heading = MARKDOWN_HEADING_PATTERN.match(lines[end])
            if next_heading and len(next_heading.group(1)) <= level:
                break
            end += 1
        return index, end
    return None


def _source_entry(line: str) -> tuple[list[str], str, str] | None:
    match = re.match(r"^\s*((?:\[\d{1,4}\]\s*)+)(.+?)\s*$", line)
    if not match:
        return None
    numbers = NUMERIC_CITATION_PATTERN.findall(match.group(1))
    payload = match.group(2).strip()
    title_link = re.search(r"\[([^\]]+)\]\((https?://[^)\s]+)\)", payload)
    if title_link:
        return numbers, title_link.group(1).strip(), title_link.group(2).strip()
    titled_url = re.match(r"(.+?)\s+(https?://\S+)$", payload)
    if titled_url:
        return numbers, titled_url.group(1).strip(), titled_url.group(2).rstrip(".,;)")
    return numbers, payload, ""


def resolve_numeric_citations(
    source: str,
    verified_sources: dict[str, dict[str, str]] | None = None,
) -> tuple[str, dict[str, int]]:
    """Turn a numbered bibliography into reader-facing title links.

    Only numeric references declared inside an explicit Sources/Источники
    section are touched, avoiding accidental removal of ordinary bracketed data.
    """
    verified_sources = verified_sources or {}
    lines = source.splitlines()
    bounds = _source_section_bounds(lines)
    report = {
        "numeric_citations_scanned": 0,
        "numeric_sources_resolved": 0,
        "numeric_unresolved": 0,
    }
    if bounds is None:
        return source, report

    start, end = bounds
    references: dict[str, dict[str, str]] = {}
    ordered: list[dict[str, str]] = []
    seen_urls: set[str] = set()
    for line in lines[start + 1:end]:
        entry = _source_entry(line)
        if entry is None:
            continue
        numbers, detected_title, detected_url = entry
        for number in numbers:
            verified = verified_sources.get(number, {})
            title = verified.get("title") or detected_title
            url = verified.get("url") or detected_url
            if title and url:
                reference = {"title": title, "url": url}
                references[number] = reference
                if url not in seen_urls:
                    ordered.append(reference)
                    seen_urls.add(url)

    body = lines[:start] + lines[end:]

    def replace(match: re.Match[str]) -> str:
        number = match.group(1)
        if number not in references and not any(
            number in (_source_entry(line) or ([], "", ""))[0]
            for line in lines[start + 1:end]
        ):
            return match.group(0)
        report["numeric_citations_scanned"] += 1
        if number in references:
            report["numeric_sources_resolved"] += 1
        else:
            report["numeric_unresolved"] += 1
        return ""

    cleaned_body = [NUMERIC_CITATION_PATTERN.sub(replace, line).rstrip() for line in body]
    while cleaned_body and not cleaned_body[-1].strip():
        cleaned_body.pop()
    sources_block: list[str] = []
    if ordered:
        heading = lines[start]
        sources_block = [heading, "", *[f"- [{entry['title']}]({entry['url']})" for entry in ordered]]
    result_lines = cleaned_body + ([""] if cleaned_body and sources_block else []) + sources_block
    return "\n".join(result_lines).strip() + "\n", report


def assert_sources_preserved(
    before: dict[str, int],
    after: dict[str, int],
    *,
    allowed_removed: set[str] | None = None,
) -> None:
    allowed_removed = allowed_removed or set()
    lost = [
        name for name, count in before.items()
        if after[name] < count and name not in allowed_removed
    ]
    if lost:
        raise ValueError(
            "Refine would remove source-bearing constructs: " + ", ".join(lost)
        )


def document_title(source: str, fallback: str) -> str:
    """Return the inline title that should become the final Obsidian filename."""
    lines = source.splitlines()
    start = 0
    if lines and lines[0].strip() == "---":
        for index in range(1, len(lines)):
            if lines[index].strip() == "---":
                start = index + 1
                break
    for line in lines[start:]:
        match = re.match(r"^#\s+(.+?)\s*$", line)
        if match:
            return normalize_text(match.group(1))
        if line.strip():
            break
    return fallback


def note_filename(source: str, fallback: str) -> str:
    title = document_title(source, fallback)
    safe_title = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "-", title)
    safe_title = re.sub(r"\s+", " ", safe_title).strip(" .")
    return f"{safe_title or fallback}.md"


def _require_reader(module: str, label: str):
    try:
        return __import__(module)
    except ImportError as error:
        raise RuntimeError(
            f"Missing {label} reader. Run Bookworm with the Codex bundled Python runtime."
        ) from error


def _write_asset(assets_dir: Path, number: int, extension: str, payload: bytes) -> Path:
    assets_dir.mkdir(parents=True, exist_ok=True)
    extension = extension.lower() if extension.lower() in IMAGE_EXTENSIONS else ".png"
    asset = assets_dir / f"figure-{number:03d}{extension}"
    asset.write_bytes(payload)
    return asset


def _markdown_table(rows: list[list[str]]) -> list[str]:
    if not rows:
        return []
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    clean = [[cell.replace("|", "\\|").replace("\n", "<br>").strip() for cell in row] for row in normalized]
    return [
        "| " + " | ".join(clean[0]) + " |",
        "| " + " | ".join("---" for _ in range(width)) + " |",
        *["| " + " | ".join(row) + " |" for row in clean[1:]],
    ]


def _docx_paragraph_markdown(paragraph) -> str:
    """Extract paragraph text while retaining external Word hyperlinks."""
    from docx.oxml.ns import qn

    parts: list[str] = []
    for child in paragraph._p:
        tag = child.tag.rsplit("}", 1)[-1]
        text = "".join(
            node.text or ""
            for node in child.iter()
            if node.tag.rsplit("}", 1)[-1] == "t"
        )
        if not text:
            continue
        if tag == "hyperlink":
            relationship_id = child.get(qn("r:id"))
            relationship = paragraph.part.rels.get(relationship_id)
            url = relationship.target_ref if relationship is not None and relationship.is_external else ""
            parts.append(f"[{text}]({url})" if url else text)
        else:
            parts.append(text)
    return "".join(parts).strip()


def convert_refine_input(source_path: Path, output_path: Path, assets_dir: Path | None = None) -> dict:
    """Convert a supported source into a temporary Markdown copy without mutation."""
    suffix = source_path.suffix.lower()
    if suffix not in REFINE_INPUT_EXTENSIONS:
        raise ValueError(f"Unsupported Refine input: {source_path.suffix or source_path}")
    if not source_path.is_file():
        raise FileNotFoundError(f"Refine input does not exist: {source_path}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    assets_dir = assets_dir or output_path.parent / "assets"
    assets: list[str] = []
    if suffix == ".md":
        markdown = source_path.read_text(encoding="utf-8")
        kind = "markdown"
    elif suffix == ".docx":
        document_module = _require_reader("docx", "python-docx")
        document = document_module.Document(source_path)
        lines: list[str] = []
        for paragraph in document.paragraphs:
            text = _docx_paragraph_markdown(paragraph)
            if not text:
                continue
            style = (paragraph.style.name or "").lower()
            heading = re.search(r"heading\s+([1-6])", style)
            if heading:
                lines.append("#" * int(heading.group(1)) + " " + text)
            elif "list" in style:
                lines.append("- " + text)
            else:
                lines.append(text)
        for table in document.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            if lines:
                lines.append("")
            lines.extend(_markdown_table(rows))
        asset_slug = slugify(document_title("\n".join(lines), source_path.stem))
        for relationship in document.part.rels.values():
            if "image" not in relationship.reltype:
                continue
            image_part = relationship.target_part
            extension = "." + image_part.content_type.rsplit("/", 1)[-1]
            asset = _write_asset(assets_dir, len(assets) + 1, extension, image_part.blob)
            assets.append(str(asset))
            lines.extend(["", f"![[assets/{asset_slug}/{asset.name}]]"])
        markdown = "\n\n".join(lines) + ("\n" if lines else "")
        kind = "docx"
    elif suffix == ".pdf":
        pdfplumber = _require_reader("pdfplumber", "pdfplumber")
        with pdfplumber.open(source_path) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        markdown = "\n\n".join(
            f"## Page {index}\n\n{text.strip()}" for index, text in enumerate(pages, start=1) if text.strip()
        )
        pypdf = _require_reader("pypdf", "pypdf")
        reader = pypdf.PdfReader(source_path)
        asset_slug = slugify(source_path.stem)
        embeds: list[str] = []
        for page in reader.pages:
            for image in page.images:
                asset = _write_asset(assets_dir, len(assets) + 1, Path(image.name).suffix, image.data)
                assets.append(str(asset))
                embeds.append(f"![[assets/{asset_slug}/{asset.name}]]")
        if embeds:
            markdown += "\n\n" + "\n\n".join(embeds)
        markdown += "\n" if markdown else ""
        kind = "pdf"
    else:
        pptx = _require_reader("pptx", "python-pptx")
        presentation = pptx.Presentation(source_path)
        sections: list[str] = []
        for index, slide in enumerate(presentation.slides, start=1):
            text = [shape.text.strip() for shape in slide.shapes if getattr(shape, "has_text_frame", False) and shape.text.strip()]
            sections.append(f"## Slide {index}\n\n" + "\n\n".join(text))
            asset_slug = slugify(source_path.stem)
            for shape in slide.shapes:
                try:
                    image = shape.image
                except (AttributeError, ValueError):
                    continue
                asset = _write_asset(assets_dir, len(assets) + 1, image.ext, image.blob)
                assets.append(str(asset))
                sections.append(f"![[assets/{asset_slug}/{asset.name}]]")
            notes = getattr(slide, "notes_slide", None)
            if notes is not None:
                note_text = [shape.text.strip() for shape in notes.shapes if getattr(shape, "has_text_frame", False) and shape.text.strip()]
                if note_text:
                    sections.append("### Notes\n\n" + "\n\n".join(note_text))
        markdown = "\n\n".join(section for section in sections if section.strip()) + "\n"
        kind = "pptx"

    if not markdown.strip():
        raise ValueError(f"No readable content extracted from {source_path}; no note was created")
    output_path.write_text(markdown, encoding="utf-8")
    return {"kind": kind, "output": str(output_path), "assets": assets}


def handoff_refined_note(
    source_path: Path,
    refined_path: Path,
    destination_dir: Path,
    *,
    confirmation: str | None,
    run_dir: Path | None = None,
    assets_dir: Path | None = None,
) -> Path:
    """Create the final note only after an explicit user confirmation token."""
    if confirmation != "user-confirmed":
        raise PermissionError("Final handoff requires explicit user confirmation")
    if not source_path.is_file() or not refined_path.is_file():
        raise FileNotFoundError("Source and refined files must exist before handoff")

    resolved_run_dir = run_dir.resolve() if run_dir is not None else None
    if resolved_run_dir is not None:
        try:
            refined_path.resolve().relative_to(resolved_run_dir)
        except ValueError as error:
            raise ValueError("Refined note must be inside the declared run directory") from error

    destination_dir.mkdir(parents=True, exist_ok=True)
    title_source = (
        source_path.read_text(encoding="utf-8")
        if source_path.suffix.lower() == ".md"
        else refined_path.read_text(encoding="utf-8")
    )
    destination = destination_dir / note_filename(title_source, source_path.stem)
    if destination.exists():
        raise FileExistsError(f"Destination already exists: {destination}")

    payload = refined_path.read_bytes()
    destination.write_bytes(payload)
    if destination.read_bytes() != payload:
        destination.unlink(missing_ok=True)
        raise OSError("Final note verification failed")

    has_assets = assets_dir is not None and assets_dir.exists() and any(
        path.is_file() for path in assets_dir.rglob("*")
    )
    if has_assets:
        assert assets_dir is not None
        asset_root = destination_dir / "assets" if destination_dir.name == "Library" else destination_dir / "assets"
        final_assets = asset_root / slugify(document_title(title_source, source_path.stem))
        if final_assets.exists():
            destination.unlink(missing_ok=True)
            raise FileExistsError(f"Asset destination already exists: {final_assets}")
        shutil.copytree(assets_dir, final_assets)
        if not any(final_assets.iterdir()):
            shutil.rmtree(final_assets)
            destination.unlink(missing_ok=True)
            raise OSError("Asset handoff verification failed")

    source_path.unlink()
    if resolved_run_dir is not None:
        shutil.rmtree(resolved_run_dir)
    elif refined_path.resolve() != source_path.resolve():
        refined_path.unlink()
    return destination


def strip_inline_title(lines: list[str]) -> list[str]:
    """Remove one document-title H1, preserving optional YAML frontmatter."""
    start = 0
    if lines and lines[0].strip() == "---":
        for index in range(1, len(lines)):
            if lines[index].strip() == "---":
                start = index + 1
                break

    for index in range(start, len(lines)):
        if not lines[index].strip():
            continue
        if re.match(r"^#\s+", lines[index]):
            return lines[:index] + lines[index + 1:]
        break
    return lines


def convert_label_value_runs(lines: list[str]) -> list[str]:
    """Turn adjacent reader-facing field/value runs into compact tables."""
    converted: list[str] = []
    index = 0
    while index < len(lines):
        match = LABEL_VALUE_PATTERN.match(lines[index])
        if not match:
            converted.append(lines[index])
            index += 1
            continue

        rows = [(match.group(1), match.group(2))]
        cursor = index + 1
        while cursor < len(lines):
            blank_start = cursor
            while cursor < len(lines) and not lines[cursor].strip():
                cursor += 1
            next_match = LABEL_VALUE_PATTERN.match(lines[cursor]) if cursor < len(lines) else None
            if next_match is None:
                cursor = blank_start
                break
            rows.append((next_match.group(1), next_match.group(2)))
            cursor += 1

        if len(rows) < 2:
            converted.append(lines[index])
            index += 1
            continue

        # Tables make compact fields scannable, but turn long research prose
        # into an unreadable narrow column. Preserve such runs as labelled
        # prose; no claim or wording is discarded.
        if any(len(value) > 240 or value.count(". ") >= 2 for _label, value in rows):
            converted.extend(lines[index:cursor])
            index = cursor
            continue

        russian = any(re.search(r"[А-Яа-яЁё]", label) for label, _value in rows)
        headers = ["Параметр", "Описание"] if russian else ["Parameter", "Description"]
        converted.extend(_markdown_table([headers, *[list(row) for row in rows]]))
        index = cursor
    return converted


def normalize_existing_tables(lines: list[str]) -> list[str]:
    """Keep existing tables readable when a link title contains a pipe."""
    normalized: list[str] = []
    in_fence = False

    def escape_link_title(match: re.Match[str]) -> str:
        title = re.sub(r"(?<!\\)\|", r"\\|", match.group(1))
        return f"[{title}]({match.group(2)})"

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```") or stripped.startswith("~~~"):
            in_fence = not in_fence
        if not in_fence and stripped.startswith("|"):
            line = re.sub(r"(?<!!)\[([^\]]+)\]\(([^)]+)\)", escape_link_title, line)
        normalized.append(line)
    return normalized


def localize_table_headers(lines: list[str]) -> list[str]:
    """Use Russian parameter headers when the research itself is Russian."""
    russian = any(re.search(r"[А-Яа-яЁё]", line) for line in lines)
    if not russian:
        return lines
    return [
        "| Параметр | Описание |" if line.strip() == "| Parameter | Description |" else line
        for line in lines
    ]


def localize_common_headings(lines: list[str]) -> list[str]:
    """Normalize generic export headings to the predominant Russian language."""
    if not any(re.search(r"[А-Яа-яЁё]", line) for line in lines):
        return lines
    replacements = {"Executive summary": "Исполнительное резюме", "Sources": "Источники"}
    result: list[str] = []
    for line in lines:
        heading = MARKDOWN_HEADING_PATTERN.match(line)
        if heading and heading.group(2).strip() in replacements:
            result.append(f"{heading.group(1)} {replacements[heading.group(2).strip()]}")
        else:
            result.append(line)
    return result


def compact_mermaid_blocks(lines: list[str]) -> list[str]:
    """Keep Mermaid editable, compact, and vertically readable as one graph."""
    compacted: list[str] = []
    index = 0
    while index < len(lines):
        if lines[index].strip().casefold() != "```mermaid":
            compacted.append(lines[index])
            index += 1
            continue

        end = index + 1
        while end < len(lines) and not lines[end].strip().startswith("```"):
            end += 1
        block = lines[index : min(end + 1, len(lines))]
        compacted.append(lines[index])
        if not any("%%{init:" in line for line in block[1:]):
            compacted.append(MERMAID_COMPACT_INIT)
        for line in lines[index + 1 : min(end + 1, len(lines))]:
            line = re.sub(r"^(\s*flowchart\s+)LR\b", r"\1TD", line, flags=re.IGNORECASE)
            compacted.append(line)
        index = end + 1
    return compacted


def remove_generated_toc(lines: list[str], toc_title: str) -> list[str]:
    """Remove the compact TOC generated by this helper before rebuilding it."""
    target = toc_title.casefold()
    for index, line in enumerate(lines):
        heading = MARKDOWN_HEADING_PATTERN.match(line)
        if not heading or len(heading.group(1)) != 2:
            continue
        if heading.group(2).strip().casefold() != target:
            continue

        end = index + 1
        has_generated_link = False
        while end < len(lines):
            next_heading = MARKDOWN_HEADING_PATTERN.match(lines[end])
            if next_heading and len(next_heading.group(1)) <= 2:
                break
            if re.match(r"^\s*- \[[^]]+\]\(#[^)]+\)\s*$", lines[end]) or re.match(
                r"^\s*- \[\[#[^|\]]+\|[^\]]+\]\]\s*$", lines[end]
            ):
                has_generated_link = True
            end += 1
        if has_generated_link:
            return lines[:index] + lines[end:]
    return lines


def main_sections(lines: list[str], toc_title: str) -> list[tuple[str, str]]:
    """Return level-two sections outside code fences for a manual TOC."""
    sections: list[tuple[str, str]] = []
    in_fence = False
    for line in lines:
        if line.strip().startswith("```") or line.strip().startswith("~~~"):
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        heading = MARKDOWN_HEADING_PATTERN.match(line)
        if not heading or len(heading.group(1)) != 2:
            continue
        title = heading.group(2).strip()
        if title.casefold() != toc_title.casefold():
            sections.append((title, title))
    return sections


def refine_markdown(source: str, toc_title: str = "Содержание") -> str:
    """Make a cosmetic Obsidian-ready copy without rewriting research content."""
    before_sources = source_counts(source)
    cleaned = CHATGPT_CITATION_PATTERN.sub("", source)
    lines = [line.rstrip() for line in cleaned.splitlines()]
    lines = strip_inline_title(lines)
    lines = remove_generated_toc(lines, toc_title)
    lines = convert_label_value_runs(lines)
    lines = normalize_existing_tables(lines)
    lines = localize_table_headers(lines)
    lines = localize_common_headings(lines)
    lines = compact_mermaid_blocks(lines)

    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()

    body = "\n".join(lines)
    sections = main_sections(lines, toc_title)
    if not sections:
        result = f"{body}\n" if body else ""
        assert_sources_preserved(
            before_sources,
            source_counts(result),
            allowed_removed={"citation_markers"},
        )
        return result

    toc = "\n".join(
        [f"## {toc_title}", ""]
        + [f"- [[#{anchor}|{title}]]" for title, anchor in sections]
    )
    result = f"{toc}\n\n{body}\n" if body else f"{toc}\n"
    assert_sources_preserved(
        before_sources,
        source_counts(result),
        allowed_removed={"citation_markers"},
    )
    return result


def parse_opf(source: BookSource) -> dict:
    names = source.names()
    opf_name = next((n for n in names if n.lower().endswith(".opf")), None)
    if not opf_name:
        raise ValueError("No OPF file found")

    root = ET.fromstring(source.read_bytes(opf_name))
    title = ""
    creators: list[str] = []
    language = ""
    manifest: dict[str, str] = {}
    spine_ids: list[str] = []

    for elem in root.iter():
        tag = elem.tag.rsplit("}", 1)[-1].lower()
        if tag == "title" and elem.text and not title:
            title = normalize_text(elem.text)
        elif tag == "creator" and elem.text:
            creators.append(normalize_text(elem.text))
        elif tag == "language" and elem.text and not language:
            language = normalize_text(elem.text)
        elif tag == "item":
            item_id = elem.attrib.get("id")
            href = elem.attrib.get("href")
            if item_id and href:
                manifest[item_id] = href
        elif tag == "itemref":
            idref = elem.attrib.get("idref")
            if idref:
                spine_ids.append(idref)

    base = str(PurePosixPath(opf_name).parent)
    if base == ".":
        base = ""

    spine: list[str] = []
    for idref in spine_ids:
        href = manifest.get(idref)
        if not href:
            continue
        spine.append(str(PurePosixPath(base) / href) if base else href)

    return {
        "opf": opf_name,
        "title": title,
        "creators": creators,
        "language": language,
        "spine": spine,
    }


def parse_html(source: BookSource, name: str) -> dict:
    raw = source.read_bytes(name).decode("utf-8", "ignore")
    parser = TextAndImageParser()
    parser.feed(raw)
    headings = [text for tag, text in parser.parts if tag.startswith("h")]
    paragraphs = [text for tag, text in parser.parts if tag in {"p", "li"}]
    fallback_headings = [
        text for text in paragraphs
        if 2 <= len(text.split()) <= 18 and len(text) <= 140
    ][:4]
    visible_headings = headings or fallback_headings
    text_words = len(normalize_text(re.sub(r"<[^>]+>", " ", raw)).split())
    return {
        "file": name,
        "title": visible_headings[0] if visible_headings else "",
        "headings": visible_headings[:12],
        "word_count": text_words,
        "image_count": len(parser.images),
        "images": parser.images,
    }


def inspect_epub(path: Path) -> dict:
    source = BookSource(path)
    opf = parse_opf(source)
    chapters = [
        parse_html(source, name)
        for name in opf["spine"]
        if name.lower().endswith(HTML_EXTENSIONS)
    ]
    image_count = sum(ch["image_count"] for ch in chapters)
    return {
        "kind": "epub",
        "path": str(path),
        "title": opf["title"],
        "creators": opf["creators"],
        "language": opf["language"],
        "chapter_count": len(chapters),
        "word_count": sum(ch["word_count"] for ch in chapters),
        "image_count": image_count,
        "chapters": chapters,
    }


def inspect_pdf(path: Path) -> dict:
    data = path.read_bytes()
    page_count = len(re.findall(rb"/Type\s*/Page\b", data))
    return {
        "kind": "pdf",
        "path": str(path),
        "page_count_estimate": page_count,
        "book_like": page_count >= 100,
    }


def inspect(path: Path) -> dict:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return inspect_pdf(path)
    if suffix == ".epub" or path.is_dir():
        return inspect_epub(path)
    raise ValueError(f"Unsupported file type: {path}")


def resolve_epub_asset(html_file: str, src: str) -> str:
    src = src.split("#", 1)[0].split("?", 1)[0]
    return str((PurePosixPath(html_file).parent / src).as_posix())


def extract_epub_assets(path: Path, out: Path, book_slug: str | None = None) -> dict:
    source = BookSource(path)
    info = parse_opf(source)
    book_slug = book_slug or slugify(info["title"] or path.stem)
    out.mkdir(parents=True, exist_ok=True)
    copied: list[dict] = []
    seen: set[str] = set()
    figure_index = 1

    for html_file in info["spine"]:
        if not html_file.lower().endswith(HTML_EXTENSIONS):
            continue
        parsed = parse_html(source, html_file)
        chapter_slug = slugify(parsed["title"] or PurePosixPath(html_file).stem, "chapter")
        chapter_out = out / chapter_slug
        for src in parsed["images"]:
            asset_name = resolve_epub_asset(html_file, src)
            if asset_name in seen:
                continue
            seen.add(asset_name)
            if not asset_name.lower().endswith(IMAGE_EXTENSIONS):
                continue
            ext = PurePosixPath(asset_name).suffix.lower() or ".png"
            dest_name = f"figure-{figure_index:03d}{ext}"
            chapter_out.mkdir(parents=True, exist_ok=True)
            dest = chapter_out / dest_name
            dest.write_bytes(source.read_bytes(asset_name))
            copied.append({
                "source": asset_name,
                "dest": str(dest),
                "obsidian_path": str(PurePosixPath("assets") / book_slug / chapter_slug / dest_name),
                "chapter": parsed["title"],
            })
            figure_index += 1

    manifest = {
        "book_slug": book_slug,
        "asset_root": str(out),
        "assets": copied,
    }
    (out / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), "utf-8")
    return manifest


def detect_vaults(roots: Iterable[Path]) -> list[dict]:
    vaults: list[dict] = []
    for root in roots:
        if not root.exists():
            continue
        for current, dirs, _files in os.walk(root):
            current_path = Path(current)
            if ".obsidian" in dirs:
                vaults.append({"path": str(current_path), "name": current_path.name})
                dirs[:] = [d for d in dirs if d != ".obsidian"]
            if len(current_path.relative_to(root).parts) >= 4:
                dirs[:] = []
    return vaults


def default_vault_roots() -> list[Path]:
    home = Path.home()
    return [home / "Desktop", home / "Documents", home]


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(description="Bookworm helper")
    sub = parser.add_subparsers(dest="command", required=True)

    inspect_cmd = sub.add_parser("inspect")
    inspect_cmd.add_argument("path", type=Path)

    vault_cmd = sub.add_parser("detect-vaults")
    vault_cmd.add_argument("--root", action="append", type=Path)

    extract_cmd = sub.add_parser("extract-epub-assets")
    extract_cmd.add_argument("path", type=Path)
    extract_cmd.add_argument("--out", required=True, type=Path)
    extract_cmd.add_argument("--book-slug")

    refine_cmd = sub.add_parser("refine-markdown")
    refine_cmd.add_argument("path", type=Path)
    refine_cmd.add_argument("--out", required=True, type=Path)
    refine_cmd.add_argument("--toc-title", default="Содержание")
    refine_cmd.add_argument("--verified-sources", type=Path)

    convert_cmd = sub.add_parser("convert-refine-input")
    convert_cmd.add_argument("path", type=Path)
    convert_cmd.add_argument("--out", required=True, type=Path)
    convert_cmd.add_argument("--assets-dir", type=Path)

    citations_cmd = sub.add_parser("inspect-citations")
    citations_cmd.add_argument("path", type=Path)

    handoff_cmd = sub.add_parser("handoff-refined-note")
    handoff_cmd.add_argument("--source", required=True, type=Path)
    handoff_cmd.add_argument("--refined", required=True, type=Path)
    handoff_cmd.add_argument("--destination-dir", required=True, type=Path)
    handoff_cmd.add_argument("--confirmation", required=True)
    handoff_cmd.add_argument("--run-dir", type=Path)
    handoff_cmd.add_argument("--assets-dir", type=Path)

    args = parser.parse_args(argv)

    if args.command == "inspect":
        print(json.dumps(inspect(args.path), ensure_ascii=False, indent=2))
    elif args.command == "detect-vaults":
        roots = args.root if args.root else default_vault_roots()
        print(json.dumps({"vaults": detect_vaults(roots)}, ensure_ascii=False, indent=2))
    elif args.command == "extract-epub-assets":
        print(json.dumps(extract_epub_assets(args.path, args.out, args.book_slug), ensure_ascii=False, indent=2))
    elif args.command == "refine-markdown":
        source = args.path.read_text(encoding="utf-8")
        verified_sources = (
            json.loads(args.verified_sources.read_text(encoding="utf-8"))
            if args.verified_sources is not None
            else {}
        )
        resolved, citation_report = resolve_citation_markers(source, verified_sources)
        resolved, numeric_report = resolve_numeric_citations(resolved, verified_sources)
        refined = refine_markdown(resolved, args.toc_title)
        source_bounds = _source_section_bounds(refined.splitlines())
        source_section = (
            "\n".join(refined.splitlines()[source_bounds[0]:source_bounds[1]])
            if source_bounds is not None
            else ""
        )
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(refined, encoding="utf-8")
        print(json.dumps({
            "input": str(args.path),
            "output": str(args.out),
            "source_counts": source_counts(source),
            "existing_title_links_preserved": source_counts(source)["markdown_links"],
            "section_sources_retained": len(MARKDOWN_SOURCE_LINK_PATTERN.findall(source_section)),
            "removed_citation_markers": source_counts(source)["citation_markers"],
            "suggested_filename": note_filename(source, args.path.stem),
            **citation_report,
            **numeric_report,
        }, ensure_ascii=False))
    elif args.command == "convert-refine-input":
        print(json.dumps(
            convert_refine_input(args.path, args.out, args.assets_dir),
            ensure_ascii=False,
        ))
    elif args.command == "inspect-citations":
        source = args.path.read_text(encoding="utf-8")
        print(json.dumps({"citations": citation_inventory(source)}, ensure_ascii=False, indent=2))
    elif args.command == "handoff-refined-note":
        destination = handoff_refined_note(
            args.source,
            args.refined,
            args.destination_dir,
            confirmation=args.confirmation,
            run_dir=args.run_dir,
            assets_dir=args.assets_dir,
        )
        print(json.dumps({"destination": str(destination)}, ensure_ascii=False))
    else:
        parser.error(f"Unknown command: {args.command}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
